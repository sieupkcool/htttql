from docx import Document

def export_exam_to_word(exam_data, file_name):
    doc = Document()
    doc.add_heading('Exam Details', level=1)

    doc.add_heading('Exam Name:', level=2)
    doc.add_paragraph(exam_data['name'])

    doc.add_heading('Subject:', level=2)
    doc.add_paragraph(exam_data['subject'])

    doc.add_heading('Date:', level=2)
    doc.add_paragraph(exam_data['date'])

    doc.add_heading('Duration:', level=2)
    doc.add_paragraph(f"{exam_data['duration']} minutes")

    doc.add_heading('Questions:', level=2)
    for question in exam_data['questions']:
        doc.add_paragraph(f"Q: {question['question']}")
        for option in question['options']:
            doc.add_paragraph(f" - {option}")

    doc.save(file_name)